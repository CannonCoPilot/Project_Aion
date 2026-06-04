#!/usr/bin/env python3
"""Build mark1-luke4-synopsis.docx from source texts."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# ----- Page margins -----
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)


def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def add_table_row(table, mark_text, luke_text):
    row = table.add_row()
    row.cells[0].text = mark_text
    row.cells[1].text = luke_text
    for cell in row.cells:
        for para in cell.paragraphs:
            para.runs[0].font.size = Pt(10) if para.runs else None


def add_synopsis_table(mark_rows, luke_rows):
    """mark_rows and luke_rows: list of (verse_ref, text) tuples."""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Mark 1'
    hdr[1].text = 'Luke 4'
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for mark_text, luke_text in zip(mark_rows, luke_rows):
        row = table.add_row()
        row.cells[0].text = mark_text
        row.cells[1].text = luke_text
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


# ===== TITLE =====
title = doc.add_heading('Gospel Synopsis: Mark 1 & Luke 4', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('King James Version · Side-by-Side Parallel Passages')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True

doc.add_paragraph('Generated 2026-04-30').runs[0].font.size = Pt(9)

doc.add_page_break()

# ===== MARK 1 FULL TEXT =====
heading('Mark Chapter 1 (KJV)', level=1)

mark_verses = [
    ("1", "¶ The beginning of the gospel of Jesus Christ, the Son of God;"),
    ("2", "As it is written in the prophets, Behold, I send my messenger before thy face, which shall prepare thy way before thee."),
    ("3", "The voice of one crying in the wilderness, Prepare ye the way of the Lord, make his paths straight."),
    ("4", "John did baptize in the wilderness, and preach the baptism of repentance for the remission of sins."),
    ("5", "And there went out unto him all the land of Judaea, and they of Jerusalem, and were all baptized of him in the river of Jordan, confessing their sins."),
    ("6", "And John was clothed with camel's hair, and with a girdle of a skin about his loins; and he did eat locusts and wild honey;"),
    ("7", "And preached, saying, There cometh one mightier than I after me, the latchet of whose shoes I am not worthy to stoop down and unloose."),
    ("8", "I indeed have baptized you with water: but he shall baptize you with the Holy Ghost."),
    ("9", "¶ And it came to pass in those days, that Jesus came from Nazareth of Galilee, and was baptized of John in Jordan."),
    ("10", "And straightway coming up out of the water, he saw the heavens opened, and the Spirit like a dove descending upon him:"),
    ("11", "And there came a voice from heaven, [saying], Thou art my beloved Son, in whom I am well pleased."),
    ("12", "And immediately the Spirit driveth him into the wilderness."),
    ("13", "And he was there in the wilderness forty days, tempted of Satan; and was with the wild beasts; and the angels ministered unto him."),
    ("14", "¶ Now after that John was put in prison, Jesus came into Galilee, preaching the gospel of the kingdom of God,"),
    ("15", "And saying, ‹The time is fulfilled, and the kingdom of God is at hand: repent ye, and believe the gospel.›"),
    ("16", "Now as he walked by the sea of Galilee, he saw Simon and Andrew his brother casting a net into the sea: for they were fishers."),
    ("17", "And Jesus said unto them, ‹Come ye after me, and I will make you to become fishers of men.›"),
    ("18", "And straightway they forsook their nets, and followed him."),
    ("19", "And when he had gone a little further thence, he saw James the [son] of Zebedee, and John his brother, who also were in the ship mending their nets."),
    ("20", "And straightway he called them: and they left their father Zebedee in the ship with the hired servants, and went after him."),
    ("21", "And they went into Capernaum; and straightway on the sabbath day he entered into the synagogue, and taught."),
    ("22", "And they were astonished at his doctrine: for he taught them as one that had authority, and not as the scribes."),
    ("23", "¶ And there was in their synagogue a man with an unclean spirit; and he cried out,"),
    ("24", "Saying, Let [us] alone; what have we to do with thee, thou Jesus of Nazareth? art thou come to destroy us? I know thee who thou art, the Holy One of God."),
    ("25", "And Jesus rebuked him, saying, ‹Hold thy peace, and come out of him.›"),
    ("26", "And when the unclean spirit had torn him, and cried with a loud voice, he came out of him."),
    ("27", "And they were all amazed, insomuch that they questioned among themselves, saying, What thing is this? what new doctrine [is] this? for with authority commandeth he even the unclean spirits, and they do obey him."),
    ("28", "And immediately his fame spread abroad throughout all the region round about Galilee."),
    ("29", "¶ And forthwith, when they were come out of the synagogue, they entered into the house of Simon and Andrew, with James and John."),
    ("30", "But Simon's wife's mother lay sick of a fever, and anon they tell him of her."),
    ("31", "And he came and took her by the hand, and lifted her up; and immediately the fever left her, and she ministered unto them."),
    ("32", "And at even, when the sun did set, they brought unto him all that were diseased, and them that were possessed with devils."),
    ("33", "And all the city was gathered together at the door."),
    ("34", "And he healed many that were sick of divers diseases, and cast out many devils; and suffered not the devils to speak, because they knew him."),
    ("35", "And in the morning, rising up a great while before day, he went out, and departed into a solitary place, and there prayed."),
    ("36", "And Simon and they that were with him followed after him."),
    ("37", "And when they had found him, they said unto him, All [men] seek for thee."),
    ("38", "And he said unto them, ‹Let us go into the next towns, that I may preach there also: for therefore came I forth.›"),
    ("39", "And he preached in their synagogues throughout all Galilee, and cast out devils."),
    ("40", "¶ And there came a leper to him, beseeching him, and kneeling down to him, and saying unto him, If thou wilt, thou canst make me clean."),
    ("41", "And Jesus, moved with compassion, put forth [his] hand, and touched him, and saith unto him, ‹I will; be thou clean.›"),
    ("42", "And as soon as he had spoken, immediately the leprosy departed from him, and he was cleansed."),
    ("43", "And he straitly charged him, and forthwith sent him away;"),
    ("44", "And saith unto him, ‹See thou say nothing to any man: but go thy way, shew thyself to the priest, and offer for thy cleansing those things which Moses commanded, for a testimony unto them.›"),
    ("45", "But he went out, and began to publish [it] much, and to blaze abroad the matter, insomuch that Jesus could no more openly enter into the city, but was without in desert places: and they came to him from every quarter."),
]

for v_num, v_text in mark_verses:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run_num = p.add_run(v_num + "  ")
    run_num.bold = True
    run_num.font.size = Pt(10)
    run_text = p.add_run(v_text)
    run_text.font.size = Pt(10)

doc.add_page_break()

# ===== LUKE 4 FULL TEXT =====
heading('Luke Chapter 4 (KJV)', level=1)

luke_verses = [
    ("1", "¶ And Jesus being full of the Holy Ghost returned from Jordan, and was led by the Spirit into the wilderness,"),
    ("2", "Being forty days tempted of the devil. And in those days he did eat nothing: and when they were ended, he afterward hungered."),
    ("3", "And the devil said unto him, If thou be the Son of God, command this stone that it be made bread."),
    ("4", "And Jesus answered him, saying, ‹It is written, That man shall not live by bread alone, but by every word of God.›"),
    ("5", "And the devil, taking him up into an high mountain, shewed unto him all the kingdoms of the world in a moment of time."),
    ("6", "And the devil said unto him, All this power will I give thee, and the glory of them: for that is delivered unto me; and to whomsoever I will I give it."),
    ("7", "If thou therefore wilt worship me, all shall be thine."),
    ("8", "And Jesus answered and said unto him, ‹Get thee behind me, Satan: for it is written, Thou shalt worship the Lord thy God, and him only shalt thou serve.›"),
    ("9", "And he brought him to Jerusalem, and set him on a pinnacle of the temple, and said unto him, If thou be the Son of God, cast thyself down from hence:"),
    ("10", "For it is written, He shall give his angels charge over thee, to keep thee:"),
    ("11", "And in [their] hands they shall bear thee up, lest at any time thou dash thy foot against a stone."),
    ("12", "And Jesus answering said unto him, ‹It is said, Thou shalt not tempt the Lord thy God.›"),
    ("13", "And when the devil had ended all the temptation, he departed from him for a season."),
    ("14", "¶ And Jesus returned in the power of the Spirit into Galilee: and there went out a fame of him through all the region round about."),
    ("15", "And he taught in their synagogues, being glorified of all."),
    ("16", "And he came to Nazareth, where he had been brought up: and, as his custom was, he went into the synagogue on the sabbath day, and stood up for to read."),
    ("17", "And there was delivered unto him the book of the prophet Esaias. And when he had opened the book, he found the place where it was written,"),
    ("18", "‹The Spirit of the Lord› [is] ‹upon me, because he hath anointed me to preach the gospel to the poor; he hath sent me to heal the brokenhearted, to preach deliverance to the captives, and recovering of sight to the blind, to set at liberty them that are bruised,›"),
    ("19", "‹To preach the acceptable year of the Lord.›"),
    ("20", "And he closed the book, and he gave [it] again to the minister, and sat down. And the eyes of all them that were in the synagogue were fastened on him."),
    ("21", "And he began to say unto them, ‹This day is this scripture fulfilled in your ears.›"),
    ("22", "And all bare him witness, and wondered at the gracious words which proceeded out of his mouth. And they said, Is not this Joseph's son?"),
    ("23", "And he said unto them, ‹Ye will surely say unto me this proverb, Physician, heal thyself: whatsoever we have heard done in Capernaum, do also here in thy country.›"),
    ("24", "And he said, ‹Verily I say unto you, No prophet is accepted in his own country.›"),
    ("25", "‹But I tell you of a truth, many widows were in Israel in the days of Elias, when the heaven was shut up three years and six months, when great famine was throughout all the land;›"),
    ("26", "‹But unto none of them was Elias sent, save unto Sarepta,› [a city] ‹of Sidon, unto a woman› [that was] ‹a widow.›"),
    ("27", "‹And many lepers were in Israel in the time of Eliseus the prophet; and none of them was cleansed, saving Naaman the Syrian.›"),
    ("28", "And all they in the synagogue, when they heard these things, were filled with wrath,"),
    ("29", "And rose up, and thrust him out of the city, and led him unto the brow of the hill whereon their city was built, that they might cast him down headlong."),
    ("30", "But he passing through the midst of them went his way,"),
    ("31", "¶ And came down to Capernaum, a city of Galilee, and taught them on the sabbath days."),
    ("32", "And they were astonished at his doctrine: for his word was with power."),
    ("33", "And in the synagogue there was a man, which had a spirit of an unclean devil, and cried out with a loud voice,"),
    ("34", "Saying, Let [us] alone; what have we to do with thee, [thou] Jesus of Nazareth? art thou come to destroy us? I know thee who thou art; the Holy One of God."),
    ("35", "And Jesus rebuked him, saying, ‹Hold thy peace, and come out of him.› And when the devil had thrown him in the midst, he came out of him, and hurt him not."),
    ("36", "And they were all amazed, and spake among themselves, saying, What a word [is] this! for with authority and power he commandeth the unclean spirits, and they come out."),
    ("37", "And the fame of him went out into every place of the country round about."),
    ("38", "And he arose out of the synagogue, and entered into Simon's house. And Simon's wife's mother was taken with a great fever; and they besought him for her."),
    ("39", "And he stood over her, and rebuked the fever; and it left her: and immediately she arose and ministered unto them."),
    ("40", "Now when the sun was setting, all they that had any sick with divers diseases brought them unto him; and he laid his hands on every one of them, and healed them."),
    ("41", "And devils also came out of many, crying out, and saying, Thou art Christ the Son of God. And he rebuking [them] suffered them not to speak: for they knew that he was Christ."),
    ("42", "And when it was day, he departed and went into a desert place: and the people sought him, and came unto him, and stayed him, that he should not depart from them."),
    ("43", "And he said unto them, ‹I must preach the kingdom of God to other cities also: for therefore am I sent.›"),
    ("44", "And he preached in the synagogues of Galilee."),
]

for v_num, v_text in luke_verses:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run_num = p.add_run(v_num + "  ")
    run_num.bold = True
    run_num.font.size = Pt(10)
    run_text = p.add_run(v_text)
    run_text.font.size = Pt(10)

doc.add_page_break()

# ===== PARALLEL SCENES =====
heading('Parallel Scenes — Side-by-Side Synopsis', level=1)

# --- Scene S1 ---
heading('Scene S1: The Temptation in the Wilderness', level=2)
note('Mark 1:12–13  ∥  Luke 4:1–13')

table1 = doc.add_table(rows=1, cols=2)
table1.style = 'Table Grid'
hdr = table1.rows[0].cells
hdr[0].text = 'Mark 1:12–13'
hdr[1].text = 'Luke 4:1–13'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

s1_rows = [
    ("12  And immediately the Spirit driveth him into the wilderness.",
     "1  ¶ And Jesus being full of the Holy Ghost returned from Jordan, and was led by the Spirit into the wilderness,"),
    ("13  And he was there in the wilderness forty days, tempted of Satan; and was with the wild beasts; and the angels ministered unto him.",
     "2  Being forty days tempted of the devil. And in those days he did eat nothing: and when they were ended, he afterward hungered."),
    ("(no parallel)",
     "3  And the devil said unto him, If thou be the Son of God, command this stone that it be made bread."),
    ("(no parallel)",
     "4  And Jesus answered him, saying, ‹It is written, That man shall not live by bread alone, but by every word of God.›"),
    ("(no parallel)",
     "5  And the devil, taking him up into an high mountain, shewed unto him all the kingdoms of the world in a moment of time."),
    ("(no parallel)",
     "6  And the devil said unto him, All this power will I give thee, and the glory of them: for that is delivered unto me; and to whomsoever I will I give it."),
    ("(no parallel)",
     "7  If thou therefore wilt worship me, all shall be thine."),
    ("(no parallel)",
     "8  And Jesus answered and said unto him, ‹Get thee behind me, Satan: for it is written, Thou shalt worship the Lord thy God, and him only shalt thou serve.›"),
    ("(no parallel)",
     "9  And he brought him to Jerusalem, and set him on a pinnacle of the temple, and said unto him, If thou be the Son of God, cast thyself down from hence:"),
    ("(no parallel)",
     "10  For it is written, He shall give his angels charge over thee, to keep thee:"),
    ("(no parallel)",
     "11  And in [their] hands they shall bear thee up, lest at any time thou dash thy foot against a stone."),
    ("(no parallel)",
     "12  And Jesus answering said unto him, ‹It is said, Thou shalt not tempt the Lord thy God.›"),
    ("13b  …and the angels ministered unto him.",
     "13  And when the devil had ended all the temptation, he departed from him for a season."),
]
for m, l in s1_rows:
    row = table1.add_row()
    row.cells[0].text = m
    row.cells[1].text = l
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

note("Note: Mark compresses the temptation into two verses; Luke expands it into three structured dialogues. Mark alone mentions the wild beasts and the angels' ministry.")
doc.add_paragraph()

# --- Scene S2 ---
heading('Scene S2: Return to Galilee / Beginning of Ministry', level=2)
note('Mark 1:14–15  ∥  Luke 4:14–15')

table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Table Grid'
hdr = table2.rows[0].cells
hdr[0].text = 'Mark 1:14–15'
hdr[1].text = 'Luke 4:14–15'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

s2_rows = [
    ("14  ¶ Now after that John was put in prison, Jesus came into Galilee, preaching the gospel of the kingdom of God,",
     "14  ¶ And Jesus returned in the power of the Spirit into Galilee: and there went out a fame of him through all the region round about."),
    ("15  And saying, ‹The time is fulfilled, and the kingdom of God is at hand: repent ye, and believe the gospel.›",
     "15  And he taught in their synagogues, being glorified of all."),
]
for m, l in s2_rows:
    row = table2.add_row()
    row.cells[0].text = m
    row.cells[1].text = l
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

note("Note: Mark notes John's imprisonment as the trigger; Luke emphasizes the Spirit's power and growing fame.")
doc.add_paragraph()

# --- Scene S3 ---
heading('Scene S3: Man with Unclean Spirit in Capernaum Synagogue', level=2)
note('Mark 1:21–28  ∥  Luke 4:31–37')

table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Table Grid'
hdr = table3.rows[0].cells
hdr[0].text = 'Mark 1:21–28'
hdr[1].text = 'Luke 4:31–37'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

s3_rows = [
    ("21  And they went into Capernaum; and straightway on the sabbath day he entered into the synagogue, and taught.",
     "31  ¶ And came down to Capernaum, a city of Galilee, and taught them on the sabbath days."),
    ("22  And they were astonished at his doctrine: for he taught them as one that had authority, and not as the scribes.",
     "32  And they were astonished at his doctrine: for his word was with power."),
    ("23  ¶ And there was in their synagogue a man with an unclean spirit; and he cried out,",
     "33  And in the synagogue there was a man, which had a spirit of an unclean devil, and cried out with a loud voice,"),
    ("24  Saying, Let [us] alone; what have we to do with thee, thou Jesus of Nazareth? art thou come to destroy us? I know thee who thou art, the Holy One of God.",
     "34  Saying, Let [us] alone; what have we to do with thee, [thou] Jesus of Nazareth? art thou come to destroy us? I know thee who thou art; the Holy One of God."),
    ("25  And Jesus rebuked him, saying, ‹Hold thy peace, and come out of him.›",
     "35  And Jesus rebuked him, saying, ‹Hold thy peace, and come out of him.› And when the devil had thrown him in the midst, he came out of him, and hurt him not."),
    ("26  And when the unclean spirit had torn him, and cried with a loud voice, he came out of him.",
     "(included in v. 35)"),
    ("27  And they were all amazed, insomuch that they questioned among themselves, saying, What thing is this? what new doctrine [is] this? for with authority commandeth he even the unclean spirits, and they do obey him.",
     "36  And they were all amazed, and spake among themselves, saying, What a word [is] this! for with authority and power he commandeth the unclean spirits, and they come out."),
    ("28  And immediately his fame spread abroad throughout all the region round about Galilee.",
     "37  And the fame of him went out into every place of the country round about."),
]
for m, l in s3_rows:
    row = table3.add_row()
    row.cells[0].text = m
    row.cells[1].text = l
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

note("Note: The demon's speech (vv. 24/34) is nearly verbatim. Luke adds that the devil 'hurt him not.' Luke attributes amazement to authority and power; Mark, to authority alone.")
doc.add_paragraph()

# --- Scene S4 ---
heading("Scene S4: Healing of Simon's Mother-in-Law", level=2)
note('Mark 1:29–31  ∥  Luke 4:38–39')

table4 = doc.add_table(rows=1, cols=2)
table4.style = 'Table Grid'
hdr = table4.rows[0].cells
hdr[0].text = 'Mark 1:29–31'
hdr[1].text = 'Luke 4:38–39'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

s4_rows = [
    ("29  ¶ And forthwith, when they were come out of the synagogue, they entered into the house of Simon and Andrew, with James and John.",
     "38  And he arose out of the synagogue, and entered into Simon's house. And Simon's wife's mother was taken with a great fever; and they besought him for her."),
    ("30  But Simon's wife's mother lay sick of a fever, and anon they tell him of her.",
     "(included in v. 38)"),
    ("31  And he came and took her by the hand, and lifted her up; and immediately the fever left her, and she ministered unto them.",
     "39  And he stood over her, and rebuked the fever; and it left her: and immediately she arose and ministered unto them."),
]
for m, l in s4_rows:
    row = table4.add_row()
    row.cells[0].text = m
    row.cells[1].text = l
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

note("Note: Mark names four disciples present; Luke simply says 'Simon's house.' Luke alone calls it 'a great fever.'")
doc.add_paragraph()

# --- Scene S5 ---
heading('Scene S5: Healing Many at Evening', level=2)
note('Mark 1:32–34  ∥  Luke 4:40–41')

table5 = doc.add_table(rows=1, cols=2)
table5.style = 'Table Grid'
hdr = table5.rows[0].cells
hdr[0].text = 'Mark 1:32–34'
hdr[1].text = 'Luke 4:40–41'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

s5_rows = [
    ("32  And at even, when the sun did set, they brought unto him all that were diseased, and them that were possessed with devils.",
     "40  Now when the sun was setting, all they that had any sick with divers diseases brought them unto him; and he laid his hands on every one of them, and healed them."),
    ("33  And all the city was gathered together at the door.",
     "(no parallel)"),
    ("34  And he healed many that were sick of divers diseases, and cast out many devils; and suffered not the devils to speak, because they knew him.",
     "41  And devils also came out of many, crying out, and saying, Thou art Christ the Son of God. And he rebuking [them] suffered them not to speak: for they knew that he was Christ."),
]
for m, l in s5_rows:
    row = table5.add_row()
    row.cells[0].text = m
    row.cells[1].text = l
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

note("Note: Luke specifies individual laying on of hands. Luke records the demons' full messianic confession. Mark alone notes the whole city gathered at the door.")
doc.add_paragraph()

# --- Scene S6 ---
heading('Scene S6: Departure and Preaching in Galilee', level=2)
note('Mark 1:35–39  ∥  Luke 4:42–44')

table6 = doc.add_table(rows=1, cols=2)
table6.style = 'Table Grid'
hdr = table6.rows[0].cells
hdr[0].text = 'Mark 1:35–39'
hdr[1].text = 'Luke 4:42–44'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

s6_rows = [
    ("35  And in the morning, rising up a great while before day, he went out, and departed into a solitary place, and there prayed.",
     "42  And when it was day, he departed and went into a desert place: and the people sought him, and came unto him, and stayed him, that he should not depart from them."),
    ("36  And Simon and they that were with him followed after him.",
     "(no parallel)"),
    ("37  And when they had found him, they said unto him, All [men] seek for thee.",
     "(merged with v. 42)"),
    ("38  And he said unto them, ‹Let us go into the next towns, that I may preach there also: for therefore came I forth.›",
     "43  And he said unto them, ‹I must preach the kingdom of God to other cities also: for therefore am I sent.›"),
    ("39  And he preached in their synagogues throughout all Galilee, and cast out devils.",
     "44  And he preached in the synagogues of Galilee."),
]
for m, l in s6_rows:
    row = table6.add_row()
    row.cells[0].text = m
    row.cells[1].text = l
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

note("Note: Mark sets the departure before dawn; Luke at daybreak. Mark alone notes the early-morning prayer and casting out devils in the summary.")
doc.add_paragraph()

doc.add_page_break()

# ===== UNIQUE PASSAGES =====
heading('Unique Passages', level=1)

heading('Mark 1 — No Luke 4 Parallel', level=2)

unique_mark = [
    ("John the Baptist's Ministry (Mark 1:1–8)", [
        ("1", "¶ The beginning of the gospel of Jesus Christ, the Son of God;"),
        ("2", "As it is written in the prophets, Behold, I send my messenger before thy face, which shall prepare thy way before thee."),
        ("3", "The voice of one crying in the wilderness, Prepare ye the way of the Lord, make his paths straight."),
        ("4", "John did baptize in the wilderness, and preach the baptism of repentance for the remission of sins."),
        ("5", "And there went out unto him all the land of Judaea, and they of Jerusalem, and were all baptized of him in the river of Jordan, confessing their sins."),
        ("6", "And John was clothed with camel's hair, and with a girdle of a skin about his loins; and he did eat locusts and wild honey;"),
        ("7", "And preached, saying, There cometh one mightier than I after me, the latchet of whose shoes I am not worthy to stoop down and unloose."),
        ("8", "I indeed have baptized you with water: but he shall baptize you with the Holy Ghost."),
    ]),
    ("The Baptism of Jesus (Mark 1:9–11)", [
        ("9", "¶ And it came to pass in those days, that Jesus came from Nazareth of Galilee, and was baptized of John in Jordan."),
        ("10", "And straightway coming up out of the water, he saw the heavens opened, and the Spirit like a dove descending upon him:"),
        ("11", "And there came a voice from heaven, [saying], Thou art my beloved Son, in whom I am well pleased."),
    ]),
    ("Call of the First Disciples (Mark 1:16–20)", [
        ("16", "Now as he walked by the sea of Galilee, he saw Simon and Andrew his brother casting a net into the sea: for they were fishers."),
        ("17", "And Jesus said unto them, ‹Come ye after me, and I will make you to become fishers of men.›"),
        ("18", "And straightway they forsook their nets, and followed him."),
        ("19", "And when he had gone a little further thence, he saw James the [son] of Zebedee, and John his brother, who also were in the ship mending their nets."),
        ("20", "And straightway he called them: and they left their father Zebedee in the ship with the hired servants, and went after him."),
    ]),
    ("Cleansing of a Leper (Mark 1:40–45)", [
        ("40", "¶ And there came a leper to him, beseeching him, and kneeling down to him, and saying unto him, If thou wilt, thou canst make me clean."),
        ("41", "And Jesus, moved with compassion, put forth [his] hand, and touched him, and saith unto him, ‹I will; be thou clean.›"),
        ("42", "And as soon as he had spoken, immediately the leprosy departed from him, and he was cleansed."),
        ("43", "And he straitly charged him, and forthwith sent him away;"),
        ("44", "And saith unto him, ‹See thou say nothing to any man: but go thy way, shew thyself to the priest, and offer for thy cleansing those things which Moses commanded, for a testimony unto them.›"),
        ("45", "But he went out, and began to publish [it] much, and to blaze abroad the matter, insomuch that Jesus could no more openly enter into the city, but was without in desert places: and they came to him from every quarter."),
    ]),
]

for section_title, verses in unique_mark:
    heading(section_title, level=3)
    for v_num, v_text in verses:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run_num = p.add_run(v_num + "  ")
        run_num.bold = True
        run_num.font.size = Pt(10)
        run_text = p.add_run(v_text)
        run_text.font.size = Pt(10)

heading('Luke 4 — No Mark 1 Parallel', level=2)

heading('Rejection at Nazareth (Luke 4:16–30)', level=3)
rejection_verses = [
    ("16", "And he came to Nazareth, where he had been brought up: and, as his custom was, he went into the synagogue on the sabbath day, and stood up for to read."),
    ("17", "And there was delivered unto him the book of the prophet Esaias. And when he had opened the book, he found the place where it was written,"),
    ("18", "‹The Spirit of the Lord› [is] ‹upon me, because he hath anointed me to preach the gospel to the poor; he hath sent me to heal the brokenhearted, to preach deliverance to the captives, and recovering of sight to the blind, to set at liberty them that are bruised,›"),
    ("19", "‹To preach the acceptable year of the Lord.›"),
    ("20", "And he closed the book, and he gave [it] again to the minister, and sat down. And the eyes of all them that were in the synagogue were fastened on him."),
    ("21", "And he began to say unto them, ‹This day is this scripture fulfilled in your ears.›"),
    ("22", "And all bare him witness, and wondered at the gracious words which proceeded out of his mouth. And they said, Is not this Joseph's son?"),
    ("23", "And he said unto them, ‹Ye will surely say unto me this proverb, Physician, heal thyself: whatsoever we have heard done in Capernaum, do also here in thy country.›"),
    ("24", "And he said, ‹Verily I say unto you, No prophet is accepted in his own country.›"),
    ("25", "‹But I tell you of a truth, many widows were in Israel in the days of Elias, when the heaven was shut up three years and six months, when great famine was throughout all the land;›"),
    ("26", "‹But unto none of them was Elias sent, save unto Sarepta,› [a city] ‹of Sidon, unto a woman› [that was] ‹a widow.›"),
    ("27", "‹And many lepers were in Israel in the time of Eliseus the prophet; and none of them was cleansed, saving Naaman the Syrian.›"),
    ("28", "And all they in the synagogue, when they heard these things, were filled with wrath,"),
    ("29", "And rose up, and thrust him out of the city, and led him unto the brow of the hill whereon their city was built, that they might cast him down headlong."),
    ("30", "But he passing through the midst of them went his way,"),
]
for v_num, v_text in rejection_verses:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run_num = p.add_run(v_num + "  ")
    run_num.bold = True
    run_num.font.size = Pt(10)
    run_text = p.add_run(v_text)
    run_text.font.size = Pt(10)

# Footer note
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run('All KJV text reproduced verbatim from source files: mark-1.txt and luke-4.txt')
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Save
out_path = os.path.join(BASE, 'mark1-luke4-synopsis.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
